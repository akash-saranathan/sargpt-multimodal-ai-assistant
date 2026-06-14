"""
Production RAG: FAISS index, OCR for scanned PDFs, SQLite-backed dedupe/manifest,
citation-rich retrieval.

FULL RAG IMPLEMENTATION WORKING: 
- FAISS index
- OCR for scanned PDFs
- SQLite-backed dedupe/manifest
- Citation-rich retrieval

EACH STEP WORKING INDIVIDUALLY LIKE SPLITTING, LOADING, EMBEDDING, INDEXING, SEARCHING: 
- SPLITTING: RecursiveCharacterTextSplitter
- LOADING: PyPDFLoader, TextLoader
- EMBEDDING: OpenAIEmbeddings
- INDEXING: FAISS.from_documents
- SEARCHING: vs.similarity_search
"""
from __future__ import annotations

import csv
import hashlib
import json
import re
import threading
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings

import api.sqlite_store as sqlite_store

_PROJECT_ROOT = Path(__file__).resolve().parent.parent
_KB_DIR = _PROJECT_ROOT / "kb"
_INDEX_DIR = _PROJECT_ROOT / "kb_faiss"

KB_RAG_EXTENSIONS = frozenset({
    ".pdf",
    ".txt",
    ".md",
    ".rst",
    ".docx",
    ".csv",
    ".json",
    ".html",
    ".htm",
})

_embeddings: OpenAIEmbeddings | None = None
_vectorstore_cache: FAISS | None = None
_cache_sig: tuple[int, float, int] | None = None
_vs_lock = threading.Lock()

# Minimum total characters from PyPDF before skipping OCR (scanned PDFs are ~empty)
_PDF_TEXT_MIN_CHARS = 40


def invalidate_kb_index_cache() -> None:
    global _vectorstore_cache, _cache_sig
    with _vs_lock:
        _vectorstore_cache = None
        _cache_sig = None


def _get_embeddings() -> OpenAIEmbeddings:
    global _embeddings
    if _embeddings is None:
        _embeddings = OpenAIEmbeddings()
    return _embeddings


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def _list_kb_doc_paths() -> list[Path]:
    _KB_DIR.mkdir(parents=True, exist_ok=True)
    return sorted(
        p
        for p in _KB_DIR.iterdir()
        if p.is_file() and p.suffix.lower() in KB_RAG_EXTENSIONS
    )


def _kb_signature() -> tuple[int, float, int] | None:
    docs = _list_kb_doc_paths()
    if not docs:
        return None
    return (
        len(docs),
        max(p.stat().st_mtime for p in docs),
        sum(p.stat().st_size for p in docs),
    )


def _compute_manifest() -> dict[str, str]:
    return {p.name: file_sha256(p) for p in _list_kb_doc_paths()}


def _html_to_plain(html: str) -> str:
    text = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_W_T = _W_NS + "t"


def _docx_text_from_ooxml(path: Path) -> str:
    parts: list[str] = []
    try:
        with zipfile.ZipFile(path, "r") as zf:
            wanted = []
            for name in zf.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                leaf = name.rsplit("/", 1)[-1]
                if leaf == "document.xml":
                    wanted.append(name)
                elif leaf.startswith("header") and leaf.endswith(".xml"):
                    wanted.append(name)
                elif leaf.startswith("footer") and leaf.endswith(".xml"):
                    wanted.append(name)
                elif name in (
                    "word/footnotes.xml",
                    "word/endnotes.xml",
                    "word/comments.xml",
                ):
                    wanted.append(name)
            for name in sorted(set(wanted)):
                try:
                    raw = zf.read(name).decode("utf-8", errors="replace")
                except KeyError:
                    continue
                try:
                    root = ET.fromstring(raw)
                except ET.ParseError:
                    continue
                for el in root.iter(_W_T):
                    if el.text:
                        parts.append(el.text)
                    if el.tail:
                        parts.append(el.tail)
    except (zipfile.BadZipFile, OSError):
        return ""
    text = "".join(parts)
    text = re.sub(r"[\t\r\f\v]+", " ", text)
    text = re.sub(r"\n+", "\n", text)
    text = re.sub(r" {2,}", " ", text).strip()
    return text


def _docx_text_python_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return ""
    try:
        d = DocxDocument(str(path))
        out: list[str] = []
        for p in d.paragraphs:
            if p.text.strip():
                out.append(p.text)
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    out.append(" | ".join(cells))
        for sec in d.sections:
            for part in (sec.header, sec.footer):
                for p in part.paragraphs:
                    if p.text.strip():
                        out.append(p.text)
        return "\n".join(out).strip()
    except Exception:
        return ""


def _normalize_citation_metadata(docs: list[Document], source_path: Path) -> list[Document]:
    fname = source_path.name
    out: list[Document] = []
    for d in docs:
        meta = dict(d.metadata)
        meta["source"] = str(source_path.resolve())
        meta["filename"] = fname
        page = meta.get("page")
        if page is None and meta.get("page_number") is not None:
            page = meta.get("page_number")
        if page is not None:
            try:
                meta["page"] = int(page)
            except (TypeError, ValueError):
                meta["page"] = page
        else:
            meta.setdefault("page", None)
        d.metadata = meta
        out.append(d)
    return out


def _load_pdf_documents(path: Path) -> list[Document]:
    meta_base = {"source": str(path.resolve()), "filename": path.name}
    loader = PyPDFLoader(str(path))
    docs = [d for d in loader.load() if d.page_content and str(d.page_content).strip()]
    total_chars = sum(len(str(d.page_content)) for d in docs)
    if total_chars >= _PDF_TEXT_MIN_CHARS:
        return _normalize_citation_metadata(docs, path)

    try:
        from langchain_community.document_loaders import UnstructuredPDFLoader

        ocr_loader = UnstructuredPDFLoader(str(path), strategy="ocr_only")
        odocs = ocr_loader.load()
        odocs = [d for d in odocs if d.page_content and str(d.page_content).strip()]
        if odocs:
            print(f"[kb] {path.name}: using OCR (PyPDF extracted {total_chars} chars).")
            return _normalize_citation_metadata(odocs, path)
    except Exception as e:
        print(f"[kb] OCR unavailable or failed for {path.name}: {e}")

    return _normalize_citation_metadata(docs, path) if docs else []


def _load_documents(path: Path) -> list[Document]:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _load_pdf_documents(path)

        if ext in {".txt", ".md", ".rst"}:
            loader = TextLoader(str(path), autodetect_encoding=True)
            docs = loader.load()
            return _normalize_citation_metadata(docs, path)

        if ext == ".docx":
            if not path.stat().st_size:
                return []
            with path.open("rb") as bf:
                hdr = bf.read(2)
            if hdr != b"PK":
                print(f"[kb] {path.name}: not a valid .docx (ZIP).")
                return []
            text = _docx_text_from_ooxml(path)
            alt = _docx_text_python_docx(path)
            if alt:
                if not text.strip():
                    text = alt
                elif alt.strip() != text.strip() and alt not in text:
                    text = f"{text}\n{alt}".strip()
            if not text.strip():
                return []
            doc = Document(page_content=text, metadata={})
            return _normalize_citation_metadata([doc], path)

        if ext == ".csv":
            lines = []
            with path.open(newline="", encoding="utf-8-sig", errors="replace") as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if i >= 8000:
                        break
                    lines.append(" | ".join(cell.strip() for cell in row))
            text = "\n".join(lines)
            if not text.strip():
                return []
            doc = Document(page_content=text, metadata={})
            return _normalize_citation_metadata([doc], path)

        if ext == ".json":
            raw = path.read_text(encoding="utf-8", errors="replace")
            try:
                data = json.loads(raw)
            except json.JSONDecodeError as e:
                print(f"[kb] Invalid JSON in {path.name}: {e}")
                return []
            text = json.dumps(data, indent=2, ensure_ascii=False)
            if len(text) > 1_500_000:
                text = text[:1_500_000] + "\n...[truncated]"
            if not text.strip():
                return []
            doc = Document(page_content=text, metadata={})
            return _normalize_citation_metadata([doc], path)

        if ext in {".html", ".htm"}:
            raw = path.read_text(encoding="utf-8", errors="replace")
            plain = _html_to_plain(raw)
            if not plain:
                return []
            doc = Document(page_content=plain, metadata={})
            return _normalize_citation_metadata([doc], path)

    except Exception as e:
        print(f"[kb] Failed to load {path.name}: {e}")
    return []


def _collect_kb_chunks_full() -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks: list[Document] = []
    seen_content_sha: set[str] = set()
    for doc_path in _list_kb_doc_paths():
        sha = file_sha256(doc_path)
        if sha in seen_content_sha:
            print(f"[kb] Skipping duplicate content: {doc_path.name} (same bytes as earlier file).")
            continue
        seen_content_sha.add(sha)
        try:
            docs = _load_documents(doc_path)
            docs = [d for d in docs if d.page_content and str(d.page_content).strip()]
            if not docs:
                continue
            for d in docs:
                d.metadata["file_sha256"] = sha
            chunks.extend(splitter.split_documents(docs))
        except Exception as e:
            print(f"[kb] Skipped {doc_path.name}: {e}")
    return chunks


def _kb_extraction_report() -> str:
    lines: list[str] = []
    for p in _list_kb_doc_paths():
        try:
            docs = _load_documents(p)
            n = sum(len(str(d.page_content or "")) for d in docs)
            if n == 0:
                extra = ""
                if p.suffix.lower() == ".pdf":
                    extra = " Try OCR dependencies (Tesseract + poppler) if scanned."
                lines.append(f"- {p.name}: 0 characters.{extra}")
            else:
                lines.append(f"- {p.name}: {n} characters (ok).")
        except Exception as e:
            lines.append(f"- {p.name}: error {e!r}")
    return "\n".join(lines) if lines else "(no files)"


def _get_vectorstore() -> FAISS | None:
    global _vectorstore_cache, _cache_sig
    sig = _kb_signature()
    if sig is None:
        with _vs_lock:
            _vectorstore_cache = None
            _cache_sig = None
        return None

    with _vs_lock:
        if _vectorstore_cache is not None and _cache_sig == sig:
            return _vectorstore_cache

        manifest_current = _compute_manifest()
        manifest_stored = sqlite_store.kb_get_manifest()
        index_faiss = _INDEX_DIR / "index.faiss"
        _, kb_max_mtime, _ = sig

        if (
            manifest_current == manifest_stored
            and manifest_current
            and index_faiss.exists()
            and index_faiss.stat().st_mtime >= kb_max_mtime
        ):
            vs = FAISS.load_local(
                str(_INDEX_DIR),
                _get_embeddings(),
                allow_dangerous_deserialization=True,
            )
            _vectorstore_cache = vs
            _cache_sig = sig
            return vs

        _INDEX_DIR.mkdir(parents=True, exist_ok=True)

        def _full_rebuild() -> FAISS | None:
            sqlite_store.kb_clear_fingerprints()
            sqlite_store.kb_clear_manifest()
            chunks = _collect_kb_chunks_full()
            if not chunks:
                _purge_index_files()
                return None
            vs = FAISS.from_documents(chunks, _get_embeddings())
            vs.save_local(str(_INDEX_DIR))
            indexed_names = {
                c.metadata.get("filename")
                for c in chunks
                if c.metadata.get("filename")
            }
            for p in _list_kb_doc_paths():
                if p.name in indexed_names:
                    sqlite_store.kb_register_fingerprint(p.name, file_sha256(p))
            sqlite_store.kb_set_manifest(_compute_manifest())
            return vs

        def _purge_index_files() -> None:
            for name in ("index.faiss", "index.pkl"):
                try:
                    (_INDEX_DIR / name).unlink(missing_ok=True)
                except OSError:
                    pass

        if not manifest_current:
            _purge_index_files()
            _vectorstore_cache = None
            _cache_sig = None
            return None

        stored_keys = set(manifest_stored)
        current_keys = set(manifest_current)
        removed = stored_keys - current_keys
        modified = {
            k
            for k in stored_keys & current_keys
            if manifest_stored[k] != manifest_current[k]
        }
        only_additions = (
            not removed
            and not modified
            and current_keys > stored_keys
            and stored_keys <= current_keys
        )

        if removed or modified or not index_faiss.exists():
            vs = _full_rebuild()
            _vectorstore_cache = vs
            _cache_sig = sig
            return vs

        if only_additions:
            new_keys = sorted(current_keys - stored_keys)
            try:
                base = FAISS.load_local(
                    str(_INDEX_DIR),
                    _get_embeddings(),
                    allow_dangerous_deserialization=True,
                )
            except Exception:
                vs = _full_rebuild()
                _vectorstore_cache = vs
                _cache_sig = sig
                return vs

            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            extra_chunks: list[Document] = []
            name_to_path = {p.name: p for p in _list_kb_doc_paths()}
            for name in new_keys:
                pth = name_to_path.get(name)
                if not pth:
                    continue
                sha = file_sha256(pth)
                if sqlite_store.kb_content_sha_already_indexed(sha):
                    sqlite_store.kb_register_fingerprint(name, sha)
                    continue
                docs = _load_documents(pth)
                docs = [d for d in docs if d.page_content and str(d.page_content).strip()]
                if not docs:
                    continue
                for d in docs:
                    d.metadata["file_sha256"] = sha
                extra_chunks.extend(splitter.split_documents(docs))
                sqlite_store.kb_register_fingerprint(name, sha)

            if extra_chunks:
                add_vs = FAISS.from_documents(extra_chunks, _get_embeddings())
                base.merge_from(add_vs)
                base.save_local(str(_INDEX_DIR))
            sqlite_store.kb_set_manifest(manifest_current)
            _vectorstore_cache = base
            _cache_sig = sig
            return base

        vs = _full_rebuild()
        _vectorstore_cache = vs
        _cache_sig = sig
        return vs


def _format_citation(meta: dict) -> str:
    fn = meta.get("filename") or Path(meta.get("source", "")).name or "unknown"
    page = meta.get("page")
    if page is not None and str(page).strip() != "":
        return f"{fn} · p.{page}"
    return fn


def search_knowledge_base_func(query: str) -> str:
    query = query.strip().strip('"').strip("'")
    if not query:
        return "No search query was provided."

    files_present = bool(_list_kb_doc_paths())
    vs = _get_vectorstore()
    if vs is None:
        if files_present:
            report = _kb_extraction_report()
            return (
                "Could not build a searchable index. Per-file status:\n"
                f"{report}\n\n"
                f"Folder: {_KB_DIR.resolve()}"
            )
        return (
            "No knowledge base documents. Add files under "
            f"'{_KB_DIR.resolve()}' (supported: {', '.join(sorted(KB_RAG_EXTENSIONS))})."
        )

    docs = vs.similarity_search(query, k=4)
    if not docs:
        return "No relevant passages were found for that query."

    parts: list[str] = []
    cite_line: list[str] = []
    for i, doc in enumerate(docs, 1):
        cit = _format_citation(doc.metadata)
        cite_line.append(f"[S{i}] {cit}")
        parts.append(f"[S{i}] ({cit})\n{doc.page_content}")

    body = "\n\n".join(parts)
    citations = "; ".join(cite_line)
    return (
        f"{body}\n\n---\nCITATIONS (use these in your answer): {citations}\n"
        "When answering, cite sources like [S1], [S2] and mention filename/page."
    )
